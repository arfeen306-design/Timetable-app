"""Date Sheet Upload — parse PDF/Excel/Word and bulk-create ExamSessions."""
from __future__ import annotations
import io
import json
import re
from typing import List, Optional

from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from pydantic import BaseModel
from sqlalchemy.orm import Session

from backend.auth.project_scope import get_project_or_404
from backend.models.base import get_db
from backend.models.project import Project, Subject
from backend.models.exam_duties_model import ExamSession

router = APIRouter()


# ── Pydantic schemas ────────────────────────────────────────────────────────

class ParsedExamRow(BaseModel):
    date_str:   str             # "2026-03-16"
    subject:    str
    start_time: str             # "09:00"
    end_time:   str             # "12:00"
    confidence: float           # 0.0–1.0
    warning:    Optional[str] = None
    raw_text:   str


class ParsedDateSheet(BaseModel):
    rows:         List[ParsedExamRow]
    total_found:  int
    warnings:     List[str]
    file_type:    str           # "excel" | "pdf" | "word"
    parser_notes: str


class ConfirmRow(BaseModel):
    date_str:   str
    subject:    str
    start_time: str
    end_time:   str


# ── Extraction helpers ──────────────────────────────────────────────────────

DATE_PATTERNS = [
    r'\b(\d{4})[\/\-](\d{1,2})[\/\-](\d{1,2})\b',                     # 2026-03-16
    r'\b(\d{1,2})[\/\-\.](\d{1,2})[\/\-\.](\d{2,4})\b',               # 16/03/2026
    r'\b(\d{1,2})\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s*(\d{2,4})\b',
    r'\b(Mon|Tue|Wed|Thu|Fri|Sat|Sun)\w*[\s,]+(\d{1,2})\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*',
]

TIME_PATTERN = re.compile(r'\b(\d{1,2}):(\d{2})\s*(AM|PM|am|pm)?\b|\b(\d{1,2})\s*(AM|PM|am|pm)\b')


def _parse_date_from_text(text: str):
    try:
        from dateutil import parser as dp
        for pat in DATE_PATTERNS:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                return dp.parse(m.group(), dayfirst=True).date()
    except Exception:
        pass
    return None


def _parse_times_from_text(text: str) -> list:
    """Return up to 2 'HH:MM' strings found in text."""
    found = []
    try:
        from dateutil import parser as dp
        for m in TIME_PATTERN.finditer(text):
            raw = m.group().strip()
            try:
                t = dp.parse(raw)
                found.append(t.strftime('%H:%M'))
                if len(found) == 2:
                    break
            except Exception:
                continue
    except ImportError:
        pass
    return found


def _looks_like_header(cells: list) -> bool:
    keywords = {'date', 'subject', 'time', 'day', 'paper', 'from', 'to',
                'start', 'end', 'exam', 'morning', 'afternoon', 'session'}
    joined = ' '.join(str(c) for c in cells).lower()
    return sum(1 for w in keywords if w in joined) >= 2


def _extract_subject(cells: list, full_text: str) -> Optional[str]:
    dt_pat = re.compile(
        r'\d{1,2}[\/\-\.]\d{1,2}|\d{1,2}:\d{2}|\b(AM|PM)\b|\b\d{4}\b',
        re.IGNORECASE,
    )
    candidates = []
    for cell in cells:
        c = str(cell).strip()
        if not c or c.isdigit() or len(c) < 3:
            continue
        if dt_pat.search(c) and len(c) < 14:
            continue
        candidates.append(c)
    if not candidates:
        return None
    return max(candidates, key=len)


def _row_to_parsed(cells: list, raw_text: str = "") -> Optional[ParsedExamRow]:
    if not cells or all(str(c).strip() == '' for c in cells):
        return None
    full = ' '.join(str(c) for c in cells)
    exam_date = _parse_date_from_text(full)
    if not exam_date:
        return None
    times = _parse_times_from_text(full)
    start = times[0] if times else None
    end   = times[1] if len(times) > 1 else None
    subj  = _extract_subject(cells, full)
    if not subj:
        return None

    confidence = 1.0
    warning: Optional[str] = None
    if not start:
        confidence = 0.5
        warning = "Could not parse start time"
    if not end:
        confidence = min(confidence, 0.7)
        warning = (warning + "; " if warning else "") + "Could not parse end time"

    return ParsedExamRow(
        date_str   = exam_date.isoformat(),
        subject    = subj,
        start_time = start or "09:00",
        end_time   = end   or "12:00",
        confidence = confidence,
        warning    = warning,
        raw_text   = (raw_text or full)[:200],
    )


# ── Parsers ─────────────────────────────────────────────────────────────────

def _parse_excel(content: bytes, ext: str) -> ParsedDateSheet:
    rows: list = []
    warnings: list = []
    try:
        if ext == 'csv':
            import csv
            text = content.decode('utf-8', errors='replace')
            raw_rows = list(csv.reader(io.StringIO(text)))
        else:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            ws = wb.active
            raw_rows = [[str(cell.value or '').strip() for cell in row]
                        for row in ws.iter_rows()]
    except Exception as e:
        return ParsedDateSheet(rows=[], total_found=0,
                               warnings=[f"Could not read file: {e}"],
                               file_type='excel', parser_notes="Parse failed.")

    for i, row in enumerate(raw_rows):
        if not any(str(c).strip() for c in row):
            continue
        if i == 0 and _looks_like_header(row):
            continue
        parsed = _row_to_parsed(row)
        if parsed:
            rows.append(parsed)
        else:
            joined = ' | '.join(str(c) for c in row[:5])
            if any(str(c).strip() for c in row[:5]):
                warnings.append(f"Row {i+1}: could not parse '{joined[:80]}'")

    return ParsedDateSheet(rows=rows, total_found=len(rows), warnings=warnings,
                           file_type='excel',
                           parser_notes=f"Read {len(raw_rows)} rows from spreadsheet.")


def _parse_pdf(content: bytes) -> ParsedDateSheet:
    rows: list = []
    warnings: list = []
    page_count = 0
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for i, row in enumerate(table):
                            if not row or not any(row):
                                continue
                            cleaned = [str(cell or '').strip() for cell in row]
                            if i == 0 and _looks_like_header(cleaned):
                                continue
                            parsed = _row_to_parsed(cleaned)
                            if parsed:
                                rows.append(parsed)
                else:
                    text = page.extract_text() or ''
                    for line in text.split('\n'):
                        line = line.strip()
                        if not line:
                            continue
                        parsed = _row_to_parsed([line], line)
                        if parsed:
                            rows.append(parsed)
    except ImportError:
        return ParsedDateSheet(rows=[], total_found=0,
                               warnings=["pdfplumber not installed on server."],
                               file_type='pdf', parser_notes="")
    except Exception as e:
        return ParsedDateSheet(rows=[], total_found=0,
                               warnings=[f"PDF parse error: {e}"],
                               file_type='pdf', parser_notes="")

    if not rows:
        warnings.append(
            "Could not extract structured data from this PDF. "
            "The PDF may be image-based (scanned). "
            "Try uploading an Excel version of the date sheet."
        )

    return ParsedDateSheet(rows=rows, total_found=len(rows), warnings=warnings,
                           file_type='pdf',
                           parser_notes=f"Extracted from {page_count} page(s).")


def _parse_word(content: bytes) -> ParsedDateSheet:
    rows: list = []
    warnings: list = []
    table_count = 0
    para_count = 0
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        table_count = len(doc.tables)
        para_count  = len(doc.paragraphs)

        for table in doc.tables:
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if i == 0 and _looks_like_header(cells):
                    continue
                parsed = _row_to_parsed(cells)
                if parsed:
                    rows.append(parsed)

        if not rows:
            for para in doc.paragraphs:
                line = para.text.strip()
                if not line:
                    continue
                parsed = _row_to_parsed([line], line)
                if parsed:
                    rows.append(parsed)
    except ImportError:
        return ParsedDateSheet(rows=[], total_found=0,
                               warnings=["python-docx not installed on server."],
                               file_type='word', parser_notes="")
    except Exception as e:
        return ParsedDateSheet(rows=[], total_found=0,
                               warnings=[f"Word parse error: {e}"],
                               file_type='word', parser_notes="")

    return ParsedDateSheet(rows=rows, total_found=len(rows), warnings=warnings,
                           file_type='word',
                           parser_notes=f"Scanned {table_count} table(s) and {para_count} paragraph(s).")


def _fuzzy_match_subject(name: str, subjects: dict) -> Optional[Subject]:
    """Return best matching Subject object or None."""
    name_lower = name.lower().strip()
    if name_lower in subjects:
        return subjects[name_lower]
    for key, subj in subjects.items():
        if key in name_lower or name_lower in key:
            return subj
    for key, subj in subjects.items():
        if len(key) >= 4 and name_lower.startswith(key[:4]):
            return subj
    return None


# ── Endpoints ───────────────────────────────────────────────────────────────

@router.post("/parse-date-sheet")
async def parse_date_sheet(
    file: UploadFile = File(...),
    project: Project = Depends(get_project_or_404),
):
    """
    Accepts PDF, Excel (.xlsx/.csv), or Word (.docx) file.
    Returns parsed rows for preview — does NOT save to DB.
    """
    content = await file.read()
    if not content:
        raise HTTPException(400, "Uploaded file is empty.")

    fname = file.filename or ''
    ext   = fname.rsplit('.', 1)[-1].lower() if '.' in fname else ''

    if ext in ('xlsx', 'xls', 'csv'):
        return _parse_excel(content, ext)
    elif ext == 'pdf':
        return _parse_pdf(content)
    elif ext in ('docx', 'doc'):
        return _parse_word(content)
    else:
        raise HTTPException(
            400,
            f"Unsupported file type '{ext}'. Upload PDF, Excel (.xlsx/.csv), or Word (.docx)."
        )


@router.post("/confirm-date-sheet")
def confirm_date_sheet(
    rows: List[ConfirmRow],
    project: Project = Depends(get_project_or_404),
    db: Session = Depends(get_db),
):
    """
    Manager reviewed the preview and clicked Confirm.
    Bulk-inserts confirmed rows as ExamSession records.
    Skips duplicates (same project + date + start_time + subject).
    """
    import datetime as _dt

    # Build subject lookup
    subjects_list = db.query(Subject).filter(Subject.project_id == project.id).all()
    subjects_by_name = {s.name.lower(): s for s in subjects_list}

    created = 0
    skipped = 0
    errors:  list = []

    for row in rows:
        try:
            exam_date  = _dt.date.fromisoformat(row.date_str)
            start_time = _dt.time.fromisoformat(row.start_time)
            end_time   = _dt.time.fromisoformat(row.end_time)
        except ValueError as e:
            errors.append(f"'{row.subject} {row.date_str}': invalid date/time — {e}")
            continue

        # Fuzzy-match subject
        subj = _fuzzy_match_subject(row.subject, subjects_by_name)
        if not subj:
            errors.append(
                f"'{row.subject}' on {row.date_str}: no matching subject found in project. "
                "Create the subject first via the Subjects page."
            )
            continue

        # Check duplicate
        existing = db.query(ExamSession).filter(
            ExamSession.project_id == project.id,
            ExamSession.date        == exam_date,
            ExamSession.start_time  == start_time,
            ExamSession.subject_id  == subj.id,
        ).first()
        if existing:
            skipped += 1
            continue

        session = ExamSession(
            project_id    = project.id,
            subject_id    = subj.id,
            date          = exam_date,
            start_time    = start_time,
            end_time      = end_time,
            room_ids_json = json.dumps([]),
        )
        db.add(session)
        created += 1

    db.commit()
    return {"created": created, "skipped": skipped, "errors": errors}
